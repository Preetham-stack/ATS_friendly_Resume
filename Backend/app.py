import os
import json
import google.generativeai as genai
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from dotenv import load_dotenv

# --- Load environment variables ---
load_dotenv()

# --- Gemini API Configuration ---
try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY not found in .env file or environment variables.")
    genai.configure(api_key=api_key)
    GEMINI_MODEL = genai.GenerativeModel('gemini-2.5-flash')  # Updated to latest stable model
except (ValueError, KeyError) as e:
    print(f"FATAL: {e}")
    GEMINI_MODEL = None
except Exception as e:
    print(f"Error configuring Gemini: {e}")
    GEMINI_MODEL = None

# --- Google Sheets Configuration ---
SHEET_SCOPE = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/drive']
SHEET_NAME = "ATSFriendlyResume_Export"

# --- Flask Configuration ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'png', 'jpg', 'jpeg'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
CORS(app)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# --- Helper Functions ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_resume_text(filepath: str) -> str:
    """Parses text from various file types (docx, pdf, images)."""
    file_extension = filepath.rsplit('.', 1)[1].lower()

    if file_extension == "docx":
        try:
            doc = Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error parsing DOCX file {filepath}: {e}")
            return ""
    elif file_extension == "pdf":
        try:
            with fitz.open(filepath) as doc:
                return "".join(page.get_text() for page in doc)
        except Exception as e:
            print(f"Error parsing PDF file {filepath}: {e}")
            return ""
    elif file_extension in {'png', 'jpg', 'jpeg'}:
        try:
            return pytesseract.image_to_string(Image.open(filepath))
        except Exception as e:
            print(f"Error performing OCR on image {filepath}: {e}")
            return ""
    return ""


def create_optimization_prompt(resume_text: str, jd_text: str, recommendations: str) -> str:
    return f'''
    You are an expert ATS Optimization AI. Analyze the provided resume and job description, then rewrite the resume
    to maximize its ATS score (target score: 85+).

    Incorporate the following specific recommendations for improvement: {recommendations}

    Return a **valid JSON only**, with this structure:
    {{
      "ats_score": <integer>,
      "optimized_resume_text": "<string>",
      "modifications_made": ["<string>", ...],
      "user_recommendations": ["<string>", ...]
    }}

    ---
    Job Description:
    {jd_text}
    ---
    Resume:
    {resume_text}
    '''


def create_analysis_prompt(resume_text: str, jd_text: str) -> str:
    return f'''
    You are an expert ATS Analyzer AI. Your task is to analyze the resume against the provided job description.
    Do NOT rewrite the resume. Instead, identify what's missing and compare the skills.

    1. Extract all technical skills, tools, and keywords from the Job Description.
    2. Extract all technical skills and tools from the Resume.
    3. Calculate a skills matching score as a percentage based on how many skills from the job description are present in the resume.

    Return a **valid JSON only**, with this structure:
    {{
      "ats_score_estimation": <integer>,
      "skills_matching_score": <integer>,
      "missing_keywords": ["<string>", ...],
      "feedback": "<string>",
      "recommendations_for_improvement": ["<string>", ...],
      "jd_skills": ["<string>", ...],
      "resume_skills": ["<string>", ...]
    }}

    ---
    Job Description:
    {jd_text}
    ---
    Resume:
    {resume_text}
    '''


def create_comprehensive_generation_prompt(
    jd_text: str = "",
    skills: str = "",
    user_details: str = "",
    existing_resume_text: str = ""
) -> str:
    return f'''
    You are a world-class resume writing AI assistant. Your task is to generate or update a resume based on the provided information.

    **Instructions:**
    1.  **Analyze Inputs:** Review all provided sections: Job Description, Skills, User Details, and Existing Resume.
    2.  **Check for Completeness:** If `user_details` are missing but other information (`jd_text`, `skills`, or `existing_resume_text`) is present, you MUST ask for them.
    3.  **Generate or Update:**
        - If `existing_resume_text` is provided, update it to align with the `jd_text` and `skills`.
        - If not, create a new resume from scratch using `jd_text`, `skills`, and `user_details`.
    4.  **Format Correctly:** The resume text must use these special format markers:
    - [H1] for the name
    - [H2] for major sections (Summary, Experience, Education)
    - [H3] for job titles or sub-sections
    - [BULLET] for bullet points
    5.  **Calculate ATS Score:** Provide an estimated ATS score for the final generated/updated resume against the job description.
    6.  **Return JSON:** Your entire output must be a single, valid JSON object.

    **Response Scenarios:**

    **A) If User Details are missing, return this JSON structure:**
    {{
      "status": "clarification_needed",
      "message": "To create a personalized resume, please provide your details. For example: name, contact number, email, LinkedIn profile, years of experience, previous job titles, and education."
    }}

    **B) If you have enough information, return this JSON structure:**
    {{
      "status": "success",
      "generated_resume_text": "<The full, formatted resume text>",
      "ats_score": <integer>,
      "extracted_skills": ["<skill1>", "<skill2>", ...]
    }}

    ---
    **Provided Information:**
    - Job Description: "{jd_text}"
    - Skills: "{skills}"
    - User Details: "{user_details}"
    - Existing Resume: "{existing_resume_text}"
    '''


def get_gemini_response(prompt: str) -> dict:
    if not GEMINI_MODEL:
        return {"error": "Gemini API is not configured."}
    try:
        response = GEMINI_MODEL.generate_content(prompt)
        text = response.text.strip()
        json_start = text.find('{')
        json_end = text.rfind('}') + 1
        if json_start == -1 or json_end == 0:
            raise ValueError("No JSON object found in the AI response.")
        cleaned_text = text[json_start:json_end]
        return json.loads(cleaned_text)
    except Exception as e:
        print(f"Error getting response from Gemini: {e}")
        problematic_text = response.text if 'response' in locals() else 'N/A'
        print(f"Problematic response text: {problematic_text}")
        return {"error": f"Failed to parse Gemini response: {e}"}


# --- API Endpoints ---
@app.route('/api/analyze', methods=['POST'])
def analyze_resume_endpoint():
    if not GEMINI_MODEL:
        return jsonify({'error': 'AI model not configured. Please check your API key.'}), 500

    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided.'}), 400

    file = request.files['resume']
    job_description = request.form.get('job_description', '')

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid or no file selected.'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    resume_text = parse_resume_text(filepath)
    if not resume_text or "parsing is not fully implemented" in resume_text:
        return jsonify({'error': 'Unable to read resume file.'}), 500

    prompt = create_analysis_prompt(resume_text, job_description)
    ai_response = get_gemini_response(prompt)

    return jsonify(ai_response)

@app.route('/api/update-resume', methods=['POST'])
def update_resume_endpoint():
    if not GEMINI_MODEL:
        return jsonify({'error': 'AI model not configured. Please check your API key.'}), 500

    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided.'}), 400

    file = request.files['resume']
    job_description = request.form.get('job_description', '')
    recommendations = request.form.get('recommendations_for_improvement', '')

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid or no file selected.'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    resume_text = parse_resume_text(filepath)
    if not resume_text or "parsing is not fully implemented" in resume_text:
        return jsonify({'error': 'Unable to read resume file.'}), 500

    prompt = create_optimization_prompt(resume_text, job_description, recommendations)
    ai_response = get_gemini_response(prompt)

    if "error" in ai_response:
        return jsonify(ai_response), 500

    optimized_filename = f"Optimized_{filename.rsplit('.', 1)[0]}.docx"
    optimized_filepath = os.path.join(app.config['UPLOAD_FOLDER'], optimized_filename)
    doc = Document()
    doc.add_paragraph(ai_response.get("optimized_resume_text", ""))
    doc.save(optimized_filepath)

    ai_response['download_path'] = f"/uploads/{optimized_filename}"
    return jsonify(ai_response)


@app.route('/api/generate-resume', methods=['POST'], strict_slashes=False)
def generate_resume_endpoint():
    if not GEMINI_MODEL:
        return jsonify({'error': 'AI model is not configured.'}), 500

    # --- Extract all possible inputs ---
    jd_text = request.form.get('job_description', '')
    skills = request.form.get('skills', '')
    user_details = request.form.get('user_details', '')
    resume_file = request.files.get('resume')
    badge_files = request.files.getlist('badges')
    image_files = request.files.getlist('images') # For profile pics or other attachments

    existing_resume_text = ""
    if resume_file and allowed_file(resume_file.filename):
        filename = secure_filename(resume_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        resume_file.save(filepath)
        existing_resume_text = parse_resume_text(filepath)

    # --- Call the new comprehensive prompt ---
    prompt = create_comprehensive_generation_prompt(
        jd_text=jd_text,
        skills=skills,
        user_details=user_details,
        existing_resume_text=existing_resume_text
    )
    ai_response = get_gemini_response(prompt)

    if "error" in ai_response:
        return jsonify(ai_response), 500

    # --- Handle conversational response ---
    if ai_response.get("status") == "clarification_needed":
        return jsonify(ai_response)

    generated_text = ai_response.get("generated_resume_text", "")
    generated_skills = ai_response.get("extracted_skills", [])

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add badges on top-right
    if badge_files:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Also consider badges from the original resume if it was an update
        for bf in badge_files: 
            if bf and allowed_file(bf.filename):
                badge_filename = secure_filename(bf.filename)
                badge_path = os.path.join(app.config['UPLOAD_FOLDER'], badge_filename)
                bf.save(badge_path)
                try:
                    p.add_run().add_picture(badge_path, height=Inches(0.6))
                    p.add_run(" ")
                except Exception as e:
                    print(f"Error adding badge {badge_filename}: {e}")

    # Parse structured resume content
    for line in generated_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('[H1]'):
            text = line.replace('[H1]', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('[H2]'):
            text = line.replace('[H2]', '').strip()
            doc.add_paragraph()  # spacing
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith('[H3]'):
            text = line.replace('[H3]', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif line.startswith('[BULLET]'):
            text = line.replace('[BULLET]', '').strip()
            doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(line)

    generated_filename = "Generated_Resume.docx"
    generated_filepath = os.path.join(app.config['UPLOAD_FOLDER'], generated_filename)
    doc.save(generated_filepath)

    # Add download path and other relevant info to the original AI response
    ai_response['download_path'] = f"/uploads/{generated_filename}"
    ai_response['content'] = generated_text # For frontend display if needed

    return jsonify(ai_response)


@app.route('/api/export-to-sheets', methods=['POST'], strict_slashes=False)
def export_to_sheets_endpoint():
    try:
        creds_path = os.path.join(BASE_DIR, 'google-sheets-credentials.json')
        creds = ServiceAccountCredentials.from_json_keyfile_name(creds_path, SHEET_SCOPE)
        client = gspread.authorize(creds)
        sheet = client.open(SHEET_NAME).sheet1
    except FileNotFoundError:
        return jsonify({"error": "Google Sheets credentials file not found."}), 500
    except Exception as e:
        return jsonify({"error": f"Google Sheets connection failed: {e}"}), 500

    data = request.get_json()
    row = [
        data.get("ats_score", ""),
        ", ".join(data.get("modifications_made", [])),
        ", ".join(data.get("user_recommendations", [])),
        data.get("optimized_resume_text", "")
    ]

    try:
        sheet.append_row(row)
        return jsonify({"success": "Data exported to Google Sheet successfully."})
    except Exception as e:
        return jsonify({"error": f"Failed to write to Google Sheet: {e}"}), 500


@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True, port=5000)
