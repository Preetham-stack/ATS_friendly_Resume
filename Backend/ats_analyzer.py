import spacy
from docx import Document

# Load the spaCy model
# Make sure to run: python -m spacy download en_core_web_sm
nlp = spacy.load("en_core_web_sm")

# --- Keywords and Rules for ATS Analysis ---
COMMON_SECTIONS = ['experience', 'education', 'skills', 'summary', 'objective', 'contact']
ACTION_VERBS = ['managed', 'led', 'developed', 'created', 'implemented', 'achieved', 'increased', 'reduced']


def analyze_resume(resume_text, job_description):
    """
    Analyzes resume text based on ATS-friendliness rules and a job description.
    Returns a score (0-100) and a list of suggestions.
    """
    score = 0
    suggestions = []

    # --- Rule 1: Presence of Standard Sections (20 points) ---
    found_sections = 0
    resume_lower = resume_text.lower()
    for section in COMMON_SECTIONS:
        if section in resume_lower:
            found_sections += 1

    if found_sections >= 4:
        score += 20
        suggestions.append("Good: Your resume includes standard sections like Experience, Education, and Skills.")
    else:
        score += found_sections * 4
        suggestions.append(
            "Improvement: Add missing standard sections (e.g., Summary, Skills, Experience) for better organization.")

    # --- Rule 2: Keyword Matching with Job Description (40 points) ---
    if job_description:
        jd_doc = nlp(job_description.lower())
        resume_doc = nlp(resume_lower)

        # Extract nouns and proper nouns as potential keywords from JD
        jd_keywords = {token.text for token in jd_doc if token.pos_ in ['NOUN', 'PROPN']}

        # Find how many keywords are in the resume
        resume_keywords = {token.text for token in resume_doc if token.pos_ in ['NOUN', 'PROPN']}

        matched_keywords = jd_keywords.intersection(resume_keywords)

        match_percentage = (len(matched_keywords) / len(jd_keywords)) * 100 if jd_keywords else 0

        if match_percentage > 70:
            score += 40
        else:
            score += int(0.40 * match_percentage)  # Prorated score

        suggestions.append(
            f"Keyword Match: You've matched {len(matched_keywords)} of {len(jd_keywords)} important keywords from the job description.")
        if match_percentage < 50:
            suggestions.append(
                "Improvement: Tailor your resume by including more keywords from the job description, especially in your experience section.")
    else:
        score += 10  # Base score if no JD is provided
        suggestions.append("Tip: Paste a job description for a more detailed analysis and a higher potential score.")

    # --- Rule 3: Use of Action Verbs (20 points) ---
    action_verb_count = 0
    for verb in ACTION_VERBS:
        if verb in resume_lower:
            action_verb_count += 1

    if action_verb_count >= 5:
        score += 20
        suggestions.append("Good: You use strong action verbs to describe your accomplishments.")
    else:
        score += action_verb_count * 4
        suggestions.append(
            "Improvement: Start your bullet points with powerful action verbs like 'Managed', 'Developed', or 'Achieved'.")

    # --- Rule 4: Measurable Metrics (20 points) ---
    # Simple check for numbers and symbols like '%' or '$'
    doc = nlp(resume_text)
    metrics_count = len([token for token in doc if token.like_num or token.text in ['%', '$']])

    if metrics_count >= 3:
        score += 20
        suggestions.append("Good: You've included measurable results and metrics (e.g., numbers, percentages).")
    else:
        score += metrics_count * 5
        suggestions.append(
            "Improvement: Quantify your achievements. Instead of 'Increased sales', try 'Increased sales by 20% in 6 months'.")

    # Ensure score does not exceed 100
    score = min(score, 100)

    # Target score check
    if score < 85:
        suggestions.append(f"Your score is {score}. Aim for 85+ by implementing the suggestions above.")

    return score, suggestions


def generate_ats_resume(text, suggestions, output_path):
    """
    Creates a simple, ATS-friendly .docx file.
    This is a basic implementation.
    """
    document = Document()
    document.add_heading('ATS Friendly Resume', 0)

    # Add the original text
    document.add_paragraph(text)

    # Add suggestions at the end
    document.add_heading('Analysis & Suggestions', level=1)
    for suggestion in suggestions:
        document.add_paragraph(suggestion, style='List Bullet')

    document.save(output_path)
